"""Stage 1 — parse uploaded files to plain text. Mature libs only; NO OCR heavy-lifting here.

Supported (multi-format, per ADR-0021 §2 / kickoff Scope 1):
  * .pdf            -> pypdf            (text-layer extraction)
  * .docx           -> python-docx      (paragraphs + tables)
  * .xlsx           -> openpyxl         (cells, sheet by sheet)
  * .csv / .tsv     -> stdlib csv       (rows)
  * .md / .txt      -> read as text

Design notes:
  * The LLM is reserved for STRUCTURED EXTRACTION (extract.py), never for parsing. Parsing is a
    deterministic library call so the AFK gate runs offline with no model.
  * Every optional heavy dep is imported lazily and degrades gracefully: if a PDF/docx/xlsx lib is
    missing, `parse_file` raises `ParseError` for that file only (the battery can skip that format)
    while text/csv/md always work on the stdlib. This keeps the offline keyword gate green even on
    a machine without the office libs installed.
  * Output is a `ParsedDoc`: normalized text + a `doc_kind` hint (resume / project / company / roster
    / roadmap) sniffed from filename + content, which the extractor uses to route heuristics.
"""
from __future__ import annotations

import codecs
import csv
import io
import re
from dataclasses import dataclass, field
from pathlib import Path


class ParseError(Exception):
    """A single file could not be parsed (unknown/again-unsupported format or missing lib)."""


class DecodeError(ParseError):
    """fixB/B1 — the bytes could not be decoded as text under ANY candidate encoding.

    A SUBCLASS of ParseError on purpose: every existing caller (`ingest_paths`, the /ingest handler,
    the AFK battery) already handles ParseError by marking that file 'failed' and surfacing the
    message, so this inherits the honest failure path without a single change upstream. The subclass
    exists so a caller that wants to say "probably an encoding problem" in its own words can tell an
    undecodable file apart from a missing-library / unsupported-format failure."""


# feat-039 (readiness §2-D): defence-in-depth against XML entity-expansion / quadratic-blowup
# ("billion laughs") bombs in docx/xlsx. IMPORTANT ATTRIBUTION (corrected): openpyxl AND python-docx
# read OOXML XML via **lxml** when it is installed (both fall back to the stdlib xml.etree only when
# lxml is absent). `defusedxml.defuse_stdlib()` patches the STDLIB stack (xml.etree, xml.sax, ...) in
# place — so it hardens ONLY the lxml-absent fallback path, NOT lxml itself. On the common (lxml
# present) path the billion-laughs protection actually comes from lxml/libxml2's own parser, which by
# default does not resolve external/parameter entities and caps internal entity expansion; a crafted
# bomb is refused there, not by defuse_stdlib. We still call defuse_stdlib() to cover the fallback
# path (and it is cheap + idempotent). The residual attack surface via lxml's remaining internal-
# entity handling is narrow, and the zip-bomb guard in guards.archive_reason is the primary
# decompression-bomb defence at the HTTP edge regardless. Best-effort: if defusedxml is absent the
# parse still works and lxml/libxml2's limits still apply.
_XML_DEFUSED = False


def _defuse_xml() -> None:
    global _XML_DEFUSED
    if _XML_DEFUSED:
        return
    try:
        import warnings
        with warnings.catch_warnings():
            # defuse_stdlib touches the deprecated cElementTree shim; the hardening still applies.
            warnings.simplefilter("ignore", DeprecationWarning)
            import defusedxml
            defusedxml.defuse_stdlib()
    except Exception:  # pragma: no cover - env without defusedxml degrades gracefully
        pass
    _XML_DEFUSED = True


# doc_kind: coarse routing hint for the extractor. Not authoritative — the extractor still inspects
# content — but lets a resume vs a project weekly vs a roster take different heuristic paths.
DOC_KINDS = ("resume", "project", "company", "roster", "roadmap", "unknown")

# EVERY PATTERN HERE WAS AN ASCII WORD UNTIL feat-049, AND THIS LIST IS A ROUTER, NOT A LABEL.
# HeuristicExtractor.extract (extract.py) branches on doc_kind and 'unknown' matches NO branch, so a
# document that sniffs to 'unknown' contributes material chunks and nothing else — no people, no
# projects. A Chinese HR system exports 「员工花名册.md」/「本周项目周报.md」/「张伟_简历.md」; not one
# contains an ASCII keyword, so the first customer's ordinary paperwork routed to 'unknown' and their
# roster ingested as zero colleagues. The content pass below did not save it either: it runs the same
# five ASCII regexes over the head of a Chinese document.
#
# THE CHINESE ALTERNATIONS MIRROR THE ENGLISH ONES RATHER THAN EXTENDING THEM. Each is the direct
# rendering of a word already on its line — 周报 for weekly, 项目 for project, 名册 for roster — so
# the router's breadth is unchanged and only its alphabet grows. Two English words are deliberately
# NOT rendered, because their Chinese equivalents are far broader than the originals:
#   * update -> 更新 is ordinary prose ("花名册由人事部于每月初更新" — a line in our own roster
#     fixture), so it would route half the corpus to 'project'.
#   * about  -> 关于 is the same problem, one kind lower.
# Likewise the roster line takes the compounds 员工名单/人员名单/团队名单 but NOT bare 名单: 名单 is
# any list of names (获奖名单, 客户名单), while 名册 is a register OF PEOPLE and nothing else.
# Simplified + Traditional throughout — a Sanya hotel takes HK/TW paperwork too. `.lower()` in
# sniff_kind is a no-op on Han, and no Han character can appear in an ASCII filename, so the English
# routes below are untouched BY CONSTRUCTION (frozen in test_sniff_kind_english_routing_is_frozen).
#
# ORDER IS LOAD-BEARING AND UNCHANGED: first match wins, so 「员工手册」-style compounds reach
# 'company' only after roster/roadmap/project have declined them, exactly as 'Company_Roster.xlsx'
# has always resolved to roster rather than company.
_KIND_HINTS: list[tuple[str, str]] = [
    (r"resume|cv|curriculum|bio\b|profile"
     r"|简历|簡歷|履历|履歷", "resume"),
    (r"roster|directory|team[_\- ]?list|headcount|people"
     r"|名册|名冊|员工名单|員工名單|人员名单|人員名單|团队名单|團隊名單|通讯录|通訊錄", "roster"),
    (r"roadmap|milestone|timeline|plan\b"
     r"|路线图|路線圖|里程碑|时间线|時間線|排期表", "roadmap"),
    (r"weekly|status|standup|update|project|sprint|retro|report|brief"
     r"|周报|週報|日报|日報|月报|月報|项目|項目|站会|站會|冲刺|衝刺|复盘|復盤|简报|簡報|汇报|匯報"
     r"|报告|報告|状态|狀態", "project"),
    (r"handbook|company|studio|org|policy|overview|about|charter|onboarding"
     r"|手册|手冊|公司|工作室|组织架构|組織架構|规章制度|規章制度|政策|概览|概覽|章程", "company"),
]


def sniff_kind(name: str, text: str) -> str:
    """Guess the document kind from filename first, then a peek at the head of the content."""
    hay_name = (name or "").lower()
    for rx, kind in _KIND_HINTS:
        if re.search(rx, hay_name):
            return kind
    head = "\n".join((text or "").splitlines()[:12]).lower()
    for rx, kind in _KIND_HINTS:
        if re.search(rx, head):
            return kind
    return "unknown"


@dataclass
class ParsedDoc:
    name: str                    # source filename (e.g. "Team_Roster.csv")
    text: str                    # normalized plain text
    doc_kind: str = "unknown"    # routing hint (see DOC_KINDS)
    ext: str = ""                # lower-cased extension without dot
    meta: dict = field(default_factory=dict)

    @property
    def lines(self) -> list[str]:
        return [ln.rstrip() for ln in self.text.splitlines()]


# feat-023: pdf text layers (pypdf) leak typographic ligatures, soft hyphens and U+FFFD
# replacement chars into the corpus — and facts.md is the advisor's citable memory, so the
# mojibake gate (test_seed_gate) demands a clean corpus. Fixed at the parse seam for every format.
#
# fixB/B1 RE-AUDIT of the U+FFFD entry. Scrubbing the replacement character kept the corpus clean
# and simultaneously destroyed the only evidence that a decode had failed: a whole GB18030 roster
# came through as a handful of Latin characters with nothing anywhere saying so. U+FFFD is therefore
# no longer a plain translate entry. It is now handled by `_audit_replacement_chars` BEFORE
# normalization, which keeps the two cases apart instead of treating them alike:
#   * a HANDFUL of them  = a broken glyph in a PDF text layer. Scrubbed, exactly as feat-023 wanted,
#     but the count is recorded on ParsedDoc.meta['replacement_chars'] so the evidence survives.
#   * a DOCUMENT FULL of them = the bytes were never decoded. Raise DecodeError; a file we could not
#     read must be reported as unread, never quietly reduced to its ASCII residue.
_MOJIBAKE_MAP = {
    "ﬀ": "ff", "ﬁ": "fi", "ﬂ": "fl", "ﬃ": "ffi", "ﬄ": "ffl",
    "ﬅ": "st", "ﬆ": "st",
    "­": "",          # soft hyphen
    "​": "", "‌": "", "‍": "", "﻿": "",   # zero-widths / BOM
    "�": "",          # U+FFFD — scrubbed ONLY after _audit_replacement_chars has ruled it sporadic
}
_MOJIBAKE_TRANS = str.maketrans(_MOJIBAKE_MAP)

REPLACEMENT_CHAR = "�"
# Thresholds for "the decode failed" vs "one bad glyph". Both must trip, so a short document with a
# single artifact (feat-023's `Ofﬁce workﬂow � broken sh­yphen`, 1 in 31 chars) stays scrubbed while
# a mojibake'd CJK document — where replacement characters are a third of the text or more — fails
# loudly. The gap between 3% and 8% is wide enough that neither case is near the line.
_FFFD_MIN_COUNT = 8
_FFFD_MIN_RATIO = 0.08


def _audit_replacement_chars(text: str, name: str) -> int:
    """Count U+FFFD in freshly-parsed text; raise DecodeError if the document is MADE of them.

    Text-family formats can no longer reach here with mojibake (`decode_text` is strict), so this
    guards the formats whose libraries do their own decoding — a PDF with a broken /ToUnicode map, an
    xlsx written by a tool that mangled its own strings. Those used to be silently emptied too."""
    n = text.count(REPLACEMENT_CHAR)
    if n >= _FFFD_MIN_COUNT and text and n / len(text) >= _FFFD_MIN_RATIO:
        label = f"{name!r}" if name else "this file"
        raise DecodeError(
            f"could not read the text in {label}: {n} of its {len(text)} characters came back as "
            f"U+FFFD (unreadable). The file's text is stored in an encoding or font map we could "
            f"not decode — a scanned/image-only PDF or a re-saved copy usually fixes it."
        )
    return n


# feat-030 P3: a NUL (0x00) is ILLEGAL in a Postgres text value — a stray one in a PDF/xlsx text
# layer would crash the DB write with an opaque error. Other C0 control chars corrupt the citable
# corpus. Scrub them all at the parse seam (every format flows through here), preserving TAB (0x09)
# and NEWLINE (0x0a) which are legitimate layout. \r is already folded to \n above.
_C0_CONTROL_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")


def _normalize(text: str) -> str:
    text = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    text = text.translate(_MOJIBAKE_TRANS)
    text = _C0_CONTROL_RE.sub("", text)   # strip NUL + other C0 controls (keep \t and \n)
    # collapse >2 blank lines, strip trailing spaces per line
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# --- text decoding (fixB/B1) --------------------------------------------------------------------
#
# THE BUG THIS REPLACES: `data.decode('utf-8', errors='replace')`. On a Chinese Windows box Excel's
# "CSV" and Notepad's "ANSI" are GB18030, not UTF-8 — the default save format for two of the three
# companies this ships to. Every Han character became U+FFFD, and `_normalize`'s _MOJIBAKE_MAP then
# mapped U+FFFD to the empty string, DESTROYING THE ONLY EVIDENCE THAT THE DECODE HAD FAILED. What
# survived was a short string of Latin garbage: doc_kind still sniffed 'roster', extraction found
# zero people, the red-line scanner found no 绩效评分 to catch (a SCORING sheet sailed through the
# gate it exists to fail), and the whole chain answered HTTP 200 / "Ingested 1 file(s)". The customer
# concludes Avery cannot read Chinese and has no way to discover otherwise.
#
# THE RULE THIS ENCODES: "I could not read it" and "the customer did not write it" are different
# facts and may never be conflated. So decoding is now STRICT — a candidate either round-trips
# cleanly or it is rejected — and when no candidate survives we raise DecodeError instead of
# inventing a plausible-looking empty document.
#
# WHY NO chardet / charset-normalizer: neither is a declared dependency (requirements.txt /
# requirements-service.txt), and adding one would make parsing depend on a statistical model that
# behaves differently in the dev venv and in the shipped image. The candidate ladder below is
# deterministic, needs nothing but the stdlib, and is exhaustively testable.
#
# WHY latin-1 IS NOT ON THE LADDER (the spec suggested it; this is a deliberate departure):
# latin-1 maps ALL 256 byte values, so it can never fail. Putting it last would mean no file is ever
# reported as undecodable — the silent-destruction bug, reintroduced one layer down. cp1252 is the
# useful part of that idea (it IS what "ANSI" means on a Western Windows box) and it CAN fail: five
# byte values (0x81/0x8D/0x8F/0x90/0x9D) are undefined, which is exactly what makes it a test rather
# than a rubber stamp.
#
# 🔴 KNOWN RESIDUAL, WRITTEN DOWN RATHER THAN GLOSSED (fixB 收口). Everything below decides between
# code pages using STRUCTURE — character blocks, script mixing, national-standard membership. That
# is decisive for every family this product actually meets, and it is NOT decisive for a single-byte
# code page we do not support whose bytes happen to land in the common-Chinese region. Concretely:
# a SHORT line of Thai (cp874) can still decode as gb18030 into ordinary-looking Chinese with a
# score of 0.000 — 'สวัสดีครับ ยินดีต้อนรับ' -> '是咽凑っ押 略勾盏橥姑押'. Longer Thai is now caught
# (the halfwidth-kana and script-mixing signals fire), short Thai is not.
# Separating "random common Chinese characters" from "Chinese" needs character-frequency or bigram
# knowledge, which is a language model, not a structural rule — so closing it properly means
# DECLARING charset-normalizer in requirements.txt / requirements-service.txt and using it as a
# veto. That is outside this change's file boundary; it is reported upward rather than half-done
# here with a threshold that would misjudge legitimate Traditional-Chinese files (measured: an
# odd-byte-run signal that catches Thai at 0.20 also flags a real Traditional roster at 0.667).

# BOM first: an explicit byte-order mark is the file telling us its encoding outright, and no
# heuristic may overrule it. ORDER IS LOAD-BEARING — the UTF-32-LE BOM (FF FE 00 00) STARTS WITH the
# UTF-16-LE BOM (FF FE), so UTF-32 must be tested first or every UTF-32-LE file is misread as UTF-16.
_BOM_ENCODINGS: tuple[tuple[bytes, str], ...] = (
    (codecs.BOM_UTF8, "utf-8-sig"),
    (codecs.BOM_UTF32_LE, "utf-32"),
    (codecs.BOM_UTF32_BE, "utf-32"),
    (codecs.BOM_UTF16_LE, "utf-16"),
    (codecs.BOM_UTF16_BE, "utf-16"),
)

# utf-8 always leads: it is what every modern exporter writes, and CJK text in any legacy code page
# is overwhelmingly unlikely to be valid UTF-8 (so a legacy file will not be stolen by this rung).
_ENC_UTF8 = "utf-8"
# Multi-byte CJK rung. gb18030 is written first (Simplified is the dominant case for all three target
# customers) but ORDER ONLY BREAKS TIES — the rung is decided by `_decode_penalty`, because these
# code pages accept each other's bytes constantly and first-match-wins produces confident nonsense:
#   * Big5 bytes are frequently ALSO valid gb18030, so a Traditional-Chinese roster from a Hong
#     Kong/Taiwan supplier would come back as wrong Han characters. A Sanya hotel takes HK/TW
#     paperwork, so that case is real rather than theoretical.
#   * 🔴 shift_jis / euc_kr WERE MISSING ENTIRELY until fixB 收口, and that was the worse hole: a
#     Japanese or Korean file has no rung of its own, gb18030 accepts its bytes, and the customer got
#     back three lines of Chinese their file never said — stored as `ingested`, chunk-counted, and
#     quotable by the advisor. A Sanya hotel receives Japanese and Korean supplier paperwork; a
#     Swedish practice receives it from Asian subcontractors. Both are now on the ladder, so the
#     TRUE reading is at least reachable, and `_decode_penalty` is what makes it win.
#
# euc_jp is deliberately NOT here. It is as permissive as gb18030 (its two-byte space is almost fully
# assigned) and it cleanly decodes ordinary Simplified-Chinese GBK bytes into plausible Japanese —
# adding it would create a NEW symmetric ambiguity in exchange for a case that does not occur:
# Windows Excel and Notepad, which is what these customers use, write Shift_JIS (CP932), never EUC-JP.
_ENC_CJK: tuple[str, ...] = ("gb18030", "big5", "shift_jis", "euc_kr")
# Western single-byte rung — Windows "ANSI" in a Latin locale (the Swedish customer's Excel).
_ENC_LATIN: tuple[str, ...] = ("cp1252",)

# ── scoring a candidate reading (lower == more believable) ───────────────────────────────────────
#
# fixB/收口 — WHY THIS GREW A SECOND AND THIRD SIGNAL. The first cut scored a candidate ONLY on the
# blocks below (`_implausibility`) and put gb18030 at the head of the CJK rung. That is enough to
# tell Simplified from Traditional, and NOT ENOUGH FOR ANYTHING ELSE, because gb18030's two-byte
# space is ~99.8% assigned: it accepts almost any byte pair and lands the result in the ordinary Han
# block, where none of these blocks fire. Measured on a Shift_JIS 「社員名簿.csv」:
#     read as gb18030 -> '巵柤,晹彁,栶怑 / 揷拞懢榊,塩嬈晹,晹挿'   implausibility 0.000
# Zero. So a Japanese roster was decoded into three lines of invented Chinese, scored perfect, and
# was stored as `ingested` with a chunk count — Avery quoting a customer sentences their file has
# never contained. That is worse than the GB18030 bug this whole change set exists to fix: that one
# lost the content, this one FABRICATES it. Same for EUC-KR Korean. Two additions close it:
#   * `_rare_han_share` — a specificity discount for the permissive codec (see below).
#   * `_native_script_share` — Hangul/kana are scripts a WRONG CJK code page does not produce.

# Blocks that correctly-decoded business paperwork essentially never contains, but that a CJK code
# page misread as the OTHER CJK code page produces constantly. The Private Use Area is the sharpest
# tell: it has no assigned meaning at all, so a decoder emitting it is guessing.
# Measured on 「姓名,職位,團隊 / 張偉,產品經理,產品組」 in Big5 (25 non-ASCII chars):
#   read as big5   -> 0 implausible characters
#   read as gb18030 -> several PUA (U+E6BD, U+E794, ...)
# and on the Simplified equivalent in GBK the verdict flips the same way.
#
# 🔴 THE BAND USED TO BE `0x2E80 <= o < 0x4E00`, WHICH PENALISED THE CORRECT ANSWER. That range
# swallows U+3000–U+303F (、。「」— ordinary CJK punctuation) and U+3040–U+30FF (kana). So the RIGHT
# reading of a Japanese sentence scored 0.800 while the fabricated-Chinese reading of the same bytes
# scored 0.000, and the same band charged a correct Traditional-Chinese decode for its 。and ：.
# Only the genuine debris blocks are penalised now: radicals/Kangxi/description characters, CJK
# strokes, and Yi.
_DEBRIS_RANGES: tuple[tuple[int, int], ...] = (
    (0x2E80, 0x2FFF),   # CJK Radicals Supplement / Kangxi Radicals / Ideographic Description
    (0x31C0, 0x31EF),   # CJK Strokes
    (0xA000, 0xA4CF),   # Yi syllables + radicals
)
_KANA_RANGES: tuple[tuple[int, int], ...] = ((0x3040, 0x30FF),)          # hiragana + katakana
_HANGUL_RANGES: tuple[tuple[int, int], ...] = ((0x1100, 0x11FF), (0x3130, 0x318F), (0xAC00, 0xD7A3))
_NATIVE_SCRIPT_RANGES: tuple[tuple[int, int], ...] = _KANA_RANGES + _HANGUL_RANGES


def _in_ranges(o: int, ranges: tuple[tuple[int, int], ...]) -> bool:
    return any(lo <= o <= hi for lo, hi in ranges)


def _implausibility(text: str) -> float:
    """Share of characters that indicate the decoder guessed wrong. 0.0 == nothing suspicious."""
    if not text:
        return 0.0
    bad = 0
    for ch in text:
        o = ord(ch)
        if 0xE000 <= o <= 0xF8FF or o >= 0xF0000:
            bad += 2          # Private Use Area — unassigned by definition
        elif _in_ranges(o, _DEBRIS_RANGES):
            bad += 1
    return bad / len(text)


# ── script coherence: the CJK scripts a reading mixes, and how much that mixing means ────────────
#
# 🔴 THE FIRST ATTEMPT AT THIS WAS A PLAIN BONUS — "any Hangul or kana is positive evidence" — AND IT
# STOLE THREE CHINESE FILES. Measured on the very corpus it was meant to protect:
#   * 「姓名,绩效评分,排名」 in GBK read as euc_kr comes out '檎츰,섀槻팀롸,탤츰' — a THIRD of it is
#     Hangul, so a flat bonus made euc_kr beat the correct gb18030 reading outright.
#   * the same roster read as big5 emits a stray 'こ', and 2 kana out of 30 characters was enough
#     bonus to beat a perfect gb18030 score.
# So presence is worthless on its own; what separates real Korean from Chinese-read-as-Korean is that
# real Korean is COHERENTLY Korean. Two gates, and the evidence only counts through both:
#   * COHERENCE. Han and kana are one family (Japanese writes 漢字 and かな in the same sentence, and
#     Chinese is Han-only — neither is "mixed"). Hangul against that family IS mixed: 「檎츰,섀槻팀롸」
#     alternates the two at ~45%, real Korean paperwork essentially never does.
#   * PRESENCE FLOOR. Real Japanese prose is 40–60% kana and real Korean 60%+ Hangul; a misread
#     sprinkles 2–6%. Below the floor the script is decoder debris, not language.
_SCRIPT_MIX_TOLERANCE = 0.10   # above this share of the minority family, the reading is incoherent
_SCRIPT_EVIDENCE_FLOOR = 0.15  # below this share, a script is a sprinkle rather than a language
_SCRIPT_EVIDENCE_WEIGHT = 0.5


def _script_profile(text: str) -> tuple[float, float]:
    """(incoherence, evidence) for a reading. `incoherence` is the share of CJK characters belonging
    to the minority script family — a decoder artefact when it is high. `evidence` is the share of
    characters written in a script a WRONG CJK code page does not produce (kana / Hangul), counted
    only when the reading is coherent and the script is actually present in quantity."""
    han = kana = hangul = 0
    for ch in text:
        o = ord(ch)
        if _in_ranges(o, _HAN_RANGES):
            han += 1
        elif _in_ranges(o, _KANA_RANGES):
            kana += 1
        elif _in_ranges(o, _HANGUL_RANGES):
            hangul += 1
    cjk = han + kana + hangul
    if not cjk or not text:
        return 0.0, 0.0
    japanese_family = han + kana
    incoherence = min(japanese_family, hangul) / cjk
    if incoherence > _SCRIPT_MIX_TOLERANCE:
        return incoherence, 0.0
    evidence = 0.0
    for count in (kana, hangul):
        share = count / len(text)
        if share >= _SCRIPT_EVIDENCE_FLOOR:
            evidence += share
    return incoherence, evidence


# cp1252 is the other codec that rubber-stamps almost anything: only five byte values are undefined,
# so a Cyrillic or Big5 file lands in it as a dense stream of accented capitals — 「姓名,職位,團隊」
# arrives as '©m¦W,Â¾¦ì,¹Î¶¤' and reads, to every downstream check, as successfully-parsed Western
# text. Genuine Swedish/German/French prose measures 5.5–8% Latin-1 supplement characters; these
# misreads measure 57–85%. Same shape of argument as `_rare_han_share`: a permissive codec has to
# show that what it produced belongs to the language it is claiming.
_LATIN_SUPPLEMENT_TOLERANCE = 0.30


def _accented_density(text: str) -> float:
    if not text:
        return 0.0
    share = sum(1 for ch in text if 0x00A0 <= ord(ch) <= 0x00FF) / len(text)
    return share if share > _LATIN_SUPPLEMENT_TOLERANCE else 0.0


# The same argument for Shift_JIS, and a hole THIS change opened: putting shift_jis on the ladder
# gave every unknown single-byte code page a new way to be "read". Shift_JIS maps 0xA1–0xDF to
# HALFWIDTH KATAKANA one byte each, so a Thai (cp874) line arrives as 'ﾊﾇﾑﾊｴﾕ､ﾃﾑｺ' — a clean decode,
# no debris blocks, no Han at all, penalty 0.000. Halfwidth katakana is a 1980s terminal/POS
# artifact: real Japanese paperwork writes prose in fullwidth kana and kanji and uses these
# essentially never, so a document MADE of them was not written in Japanese.
_HALFWIDTH_KANA_TOLERANCE = 0.30


def _halfwidth_kana_density(text: str) -> float:
    if not text:
        return 0.0
    share = sum(1 for ch in text if 0xFF61 <= ord(ch) <= 0xFF9F) / len(text)
    return share if share > _HALFWIDTH_KANA_TOLERANCE else 0.0


# The specificity discount for gb18030. A national standard's COMMON CORE is a real, checkable thing:
# GB2312 is the 6763 characters mainland paperwork is written in, Big5 the 13 060 characters
# Traditional paperwork is written in. A character in NEITHER is not something a Chinese HR export
# contains — it is what falls out of a decoder that was handed somebody else's code page. Measured
# (share of Han characters in neither GB2312 nor Big5):
#     Simplified roster in GBK ............... 0.000     Traditional roster in Big5 ...... 0.000
#     Traditional prose in Big5 .............. 0.000     Simplified roster w/ rare names . 0.056
#     Japanese roster READ AS gb18030 ........ 0.417     Japanese prose READ AS gb18030 .. 0.483
#     Traditional roster READ AS gb18030 ..... 0.167
# Every legitimate reading is at/near zero and every cross-code-page misreading is an order of
# magnitude above them, which is what makes this a usable signal rather than a guess.
#
# 🔴 IT IS DELIBERATELY ASYMMETRIC — applied to gb18030 and to nothing else. This is not "score the
# Chineseness of the text", it is "gb18030 accepts nearly every byte pair on earth, so when it claims
# a file is Chinese it has to show that what it produced IS Chinese." Charging the same fee to big5 /
# shift_jis / euc_kr would be incoherent (their output is Traditional / Japanese / Korean BY
# CONSTRUCTION) and would penalise the very readings this exists to let win.
_PERMISSIVE_ENCODINGS = frozenset({"gb18030"})
_COMMON_CORE_CODECS = ("gb2312", "big5")
_HAN_RANGES: tuple[tuple[int, int], ...] = ((0x3400, 0x9FFF), (0xF900, 0xFAFF))


def _rare_han_share(text: str) -> float:
    """Share of Han characters that are in NO Chinese national standard's common core."""
    han = [ch for ch in text if _in_ranges(ord(ch), _HAN_RANGES)]
    if not han:
        return 0.0
    rare = 0
    for ch in han:
        for codec in _COMMON_CORE_CODECS:
            try:
                ch.encode(codec)
                break
            except (UnicodeEncodeError, LookupError):
                continue
        else:
            rare += 1
    return rare / len(han)


# The small-scale version of the same lie. `b"caf\x81e"` is not valid utf-8 and 0x81 is undefined in
# cp1252, so the bytes reach gb18030, which reads 81 65 as 乪 and hands back 'caf乪' — one invented
# Han character wedged between Latin letters, ingested without a murmur. Real CJK text comes in runs;
# a lone ideograph glued to ASCII letters on either side is a decoder artefact, not a word.
def _isolated_cjk_share(text: str) -> float:
    cjk = [
        i for i, ch in enumerate(text)
        if _in_ranges(ord(ch), _HAN_RANGES) or _in_ranges(ord(ch), _NATIVE_SCRIPT_RANGES)
    ]
    if not cjk:
        return 0.0

    def is_cjk(i: int) -> bool:
        if i < 0 or i >= len(text):
            return False
        o = ord(text[i])
        return _in_ranges(o, _HAN_RANGES) or _in_ranges(o, _NATIVE_SCRIPT_RANGES)

    def is_ascii_word(i: int) -> bool:
        return 0 <= i < len(text) and text[i].isascii() and text[i].isalnum()

    stranded = sum(
        1 for i in cjk
        if not is_cjk(i - 1) and not is_cjk(i + 1) and (is_ascii_word(i - 1) or is_ascii_word(i + 1))
    )
    return stranded / len(cjk)


def _decode_penalty(text: str, enc: str) -> float:
    """How much this reading looks like a decoder guessing. Lower is better; may go NEGATIVE when a
    reading carries positive evidence (coherent kana/Hangul) no rival code page could have produced."""
    incoherence, evidence = _script_profile(text)
    penalty = _implausibility(text) + _isolated_cjk_share(text) + incoherence
    if enc in _PERMISSIVE_ENCODINGS:
        penalty += _rare_han_share(text)
    elif enc in _ENC_LATIN:
        penalty += _accented_density(text)
    elif enc == "shift_jis":
        penalty += _halfwidth_kana_density(text)
    return penalty - _SCRIPT_EVIDENCE_WEIGHT * evidence


# A reading at or under this is believable enough to stop looking (later rungs are not consulted).
_DECODE_GOOD_ENOUGH = 0.15
# A reading above this is NOT offered to the user at all. There is no third option here: shipping it
# means Avery inventing content, and pretending the file was empty means Avery inventing an absence.
# Both are the thing this change set exists to stop, so the file is reported as unread.
_DECODE_REFUSE = 0.65


def _looks_multibyte(data: bytes) -> bool:
    """Do these raw bytes look like a multi-byte CJK code page rather than a Western single-byte one?
    This orders the ladder, and it reads the RAW BYTES, so it cannot be fooled by a decode that has
    already gone wrong.

    TWO signals, either sufficient, because each one alone has a documented blind spot:

      * RUN STRUCTURE. In a CJK code page, characters cost >= 2 consecutive high bytes and Han text
        runs several characters at a time; in cp1252 Swedish/German prose the accents sit INSIDE
        otherwise-ASCII words ('Björn', 'Malmö') so the runs are length 1. This is what stops
        'Björn' (42 6A F6 72 6E) from decoding CLEANLY as gb18030 — 0xF6 is a lead byte, 0x72 a valid
        trail byte — and turning a Swedish surname into a stray Han character.
      * 🔴 HIGH-BYTE DENSITY, added at fixB 收口 because run structure ALONE MISCLASSIFIED BIG5.
        Big5 and Shift_JIS trail bytes may be in the ASCII range (0x40–0x7E): 姓 is A9 6D, 名 is A6 57.
        So 「姓名,職位,團隊」 has NO high-byte run longer than 1 and scored exactly like Swedish —
        the Traditional-Chinese ladder the comments promised was never actually reached, and the file
        went to cp1252 as '©m¦W,Â¾¦ì,¹Î¶¤'. Density does not care where the trail byte sits:
        measured, CJK files are 57–88% high bytes and Western ones 5.5–7.9%.

    Density alone is not sufficient either — an English CSV with a few Chinese names is ~20% high
    bytes — which is why run structure stays. Between them the observed gap is wide in both
    directions."""
    runs: list[int] = []
    run = 0
    for b in data:
        if b >= 0x80:
            run += 1
        elif run:
            runs.append(run)
            run = 0
    if run:
        runs.append(run)
    total = sum(runs)
    if total == 0:
        return False   # pure ASCII: the families are indistinguishable and it does not matter
    if total / len(data) >= 0.25:
        return True
    return sum(r for r in runs if r >= 2) / total >= 0.7


def _candidate_ladder(data: bytes) -> tuple[tuple[str, ...], ...]:
    """The rungs for this payload, best guess first. Encodings sharing a rung are siblings decided on
    `_decode_penalty`, not on order. (utf-8 is not a rung — `decode_text` settles it first.)

    🔴 WESTERN-SHAPED BYTES ARE NEVER READ AS CJK. The first cut kept the CJK rung as a last-resort
    fallback for them, which is how `b"caf\\x81e"` — four Latin letters and one byte cp1252 does not
    define — came back as 'caf乪', an invented ideograph stapled to an English word and ingested
    without a murmur. There is no file for which "the bytes look like isolated Western accents, and
    no Western code page accepts them, therefore it is Chinese" is sound reasoning. When the Latin
    rung refuses such a file the honest answer is that we could not read it, so the ladder ends."""
    if _looks_multibyte(data):
        return (_ENC_CJK, _ENC_LATIN)
    return (_ENC_LATIN,)


def _decode_rung(data: bytes, rung: tuple[str, ...]) -> tuple[str, str, float] | None:
    """Decode with every encoding on one rung and return (text, encoding, penalty) for the most
    believable reading, or None if the bytes are valid under none of them. Ties keep the rung's
    declared order (`min` returns the first minimal element), so gb18030 still wins whenever Big5
    offers nothing better."""
    results: list[tuple[float, str, str]] = []
    for enc in rung:
        try:
            text = data.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
        results.append((_decode_penalty(text, enc), enc, text))
    if not results:
        return None
    penalty, enc, text = min(results, key=lambda r: r[0])
    return text, enc, penalty


def decode_text(data: bytes, name: str = "") -> tuple[str, str, float]:
    """Decode uploaded bytes to text. Returns (text, encoding_used, penalty); raises DecodeError if
    nothing believable fits. STRICT throughout — this never returns a string containing damage it
    invented, and never returns a reading it does not actually believe.

    THE RUNGS ARE A PRIOR, NOT A VERDICT (fixB 收口). The first cut returned the first rung that
    decoded at all, which made the whole result hostage to `_looks_multibyte` guessing right: one
    misclassified file and a plausible-looking wrong answer was final. Now a rung short-circuits only
    when its best reading is actually believable (`_DECODE_GOOD_ENOUGH`); otherwise the search
    continues and the globally best reading wins. A German file mis-routed to the CJK rung therefore
    still comes back as German."""
    if not data:
        return "", _ENC_UTF8, 0.0
    for bom, enc in _BOM_ENCODINGS:
        if data.startswith(bom):
            try:
                return data.decode(enc), enc, 0.0
            except UnicodeDecodeError:
                # A BOM whose body does not decode is a truncated/corrupt file, not a hint worth
                # trusting. Fall through to the ladder rather than fail outright.
                break
    # utf-8 wins outright when it is valid: it is what every modern exporter writes, and a legacy CJK
    # file being accidentally valid utf-8 beyond a character or two does not happen. Deciding it here
    # rather than by score also means no scoring change can ever regress an ordinary UTF-8 upload.
    try:
        return data.decode(_ENC_UTF8), _ENC_UTF8, 0.0
    except UnicodeDecodeError:
        pass
    tried: list[str] = [_ENC_UTF8]
    best: tuple[str, str, float] | None = None
    for rung in _candidate_ladder(data):
        tried.extend(rung)
        won = _decode_rung(data, rung)
        if won is not None and (best is None or won[2] < best[2]):
            best = won
        if best is not None and best[2] <= _DECODE_GOOD_ENOUGH:
            break
    label = f"{name!r}" if name else "this file"
    if best is None:
        raise DecodeError(
            f"could not read {label} as text: its bytes are not valid in any encoding we try "
            f"({', '.join(tried)}). It is most likely saved in a different character encoding — "
            f"re-save it as UTF-8 (in Excel: 'CSV UTF-8') and upload it again."
        )
    if best[2] > _DECODE_REFUSE:
        raise DecodeError(
            f"could not work out what character encoding {label} is written in. The closest match "
            f"({best[1]}) turns it into text that does not read as any real language, so it has NOT "
            f"been read — showing you those characters would mean putting words in your file's "
            f"mouth. Re-save it as UTF-8 (in Excel: 'CSV UTF-8') and upload it again."
        )
    return best


# --- per-format extractors --------------------------------------------------------------------

def _parse_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except Exception as e:  # pragma: no cover - env without pypdf
        raise ParseError(f"pypdf not available for PDF parse: {e}")
    reader = PdfReader(io.BytesIO(data))
    # feat-039 (readiness §2-D): cap page count before extracting — a pathological PDF with a huge /
    # cyclic page tree is a CPU/RAM DoS. `len(reader.pages)` reads the (cheap) page tree; refuse over
    # the cap rather than iterate a hostile document. (A true wall-clock timeout is not portable to
    # Windows dev; the page cap is the deterministic, cross-platform guard.)
    from . import guards
    n_pages = len(reader.pages)
    if n_pages > guards.max_pdf_pages():
        raise ParseError(f"PDF has {n_pages} pages (limit {guards.max_pdf_pages()}) — refused")
    return "\n\n".join((page.extract_text() or "") for page in reader.pages)


def _parse_docx(data: bytes) -> str:
    try:
        import docx  # python-docx
    except Exception as e:  # pragma: no cover - env without python-docx
        raise ParseError(f"python-docx not available for docx parse: {e}")
    _defuse_xml()
    doc = docx.Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _parse_xlsx(data: bytes) -> str:
    try:
        from openpyxl import load_workbook
    except Exception as e:  # pragma: no cover - env without openpyxl
        raise ParseError(f"openpyxl not available for xlsx parse: {e}")
    _defuse_xml()
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    out: list[str] = []
    for ws in wb.worksheets:
        out.append(f"# sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = ["" if v is None else str(v) for v in row]
            if any(c.strip() for c in cells):
                out.append(" | ".join(cells))
    return "\n".join(out)


# The text-family parsers return the detected encoding AND its penalty alongside the text, so
# `parse_bytes` can put both on ParsedDoc.meta. Knowing a roster was read as gb18030 rather than
# utf-8 — and how confident that reading was — is the difference between diagnosing a support ticket
# in a minute and guessing at it.

def _parse_csv(data: bytes, name: str = "", delimiter: str = ",") -> tuple[str, str, float]:
    text, enc, penalty = decode_text(data, name)
    rows = list(csv.reader(io.StringIO(text), delimiter=delimiter))
    joined = "\n".join(" | ".join(cell.strip() for cell in row) for row in rows if any(row))
    return joined, enc, penalty


def _parse_text(data: bytes, name: str = "") -> tuple[str, str, float]:
    return decode_text(data, name)


# Binary/OOXML formats: the library owns the decode, so there is no encoding of ours to report.
_DISPATCH = {
    "pdf": lambda d, n="": (_parse_pdf(d), "", 0.0),
    "docx": lambda d, n="": (_parse_docx(d), "", 0.0),
    "xlsx": lambda d, n="": (_parse_xlsx(d), "", 0.0),
    "csv": lambda d, n="": _parse_csv(d, n, ","),
    "tsv": lambda d, n="": _parse_csv(d, n, "\t"),
    "md": _parse_text,
    "markdown": _parse_text,
    "txt": _parse_text,
    "text": _parse_text,
    "": _parse_text,  # extension-less -> treat as text
}


def parse_bytes(name: str, data: bytes, *, ext: str | None = None) -> ParsedDoc:
    """Parse an in-memory upload (name + raw bytes) into a ParsedDoc. Used by the HTTP upload path
    and by tests that don't want to touch disk.

    fixB/B1: raises DecodeError (a ParseError) rather than returning a plausible-looking husk when
    the bytes cannot be decoded. Callers already treat ParseError as 'this file failed', which is
    the truth we owe the user."""
    ext = (ext or Path(name).suffix.lstrip(".")).lower()
    fn = _DISPATCH.get(ext)
    if fn is None:
        raise ParseError(f"unsupported file type '.{ext}' for {name!r}")
    raw, encoding, penalty = fn(data, name)
    n_replacement = _audit_replacement_chars(raw, name)
    text = _normalize(raw)
    meta: dict = {"bytes": len(data)}
    if encoding:
        meta["encoding"] = encoding
        # fixB 收口 — 'encoding' alone reads as a fact ({'encoding': 'gb18030'} looked authoritative
        # on the Japanese roster it had just invented). Ship the confidence WITH the claim: 'high'
        # means the reading carried positive evidence for itself, 'low' means it was merely the
        # least-bad option that cleared the refuse bar. Support and the AFK battery can both see it.
        meta["decode_confidence"] = "high" if penalty <= _DECODE_GOOD_ENOUGH else "low"
        meta["decode_penalty"] = round(penalty, 3)
    if n_replacement:
        # Evidence, not decoration: _normalize is about to delete these characters, and something
        # has to remember they were there.
        meta["replacement_chars"] = n_replacement
    return ParsedDoc(name=name, text=text, doc_kind=sniff_kind(name, text), ext=ext, meta=meta)


def parse_file(path: str | Path) -> ParsedDoc:
    """Parse a file on disk into a ParsedDoc."""
    path = Path(path)
    if not path.exists():
        raise ParseError(f"no such file: {path}")
    return parse_bytes(path.name, path.read_bytes(), ext=path.suffix.lstrip("."))
